from docx import Document
from copy import deepcopy
import os

def make_word(db_manager, self):
    date_start = self.date_start.date().toString('yyyy-MM-dd')
    date_end = self.date_end.date().toString('yyyy-MM-dd')
    print(date_end, date_start)

    def fetch_data(db_manager):
        query_data = f"SELECT `date`, chk_id, dn1_id, dn2_id, dng_id, chnk27_id, pchnk271_id, pchnk272_id, pchnk273_id, schp_id FROM data_table WHERE date BETWEEN '{date_start}' AND '{date_end}'"
        query_kurs = "SELECT id, pib FROM kurs"

        data_table = db_manager.execute_query(query_data)
        kurs_table = db_manager.execute_query(query_kurs)

        return data_table, kurs_table

    data_table, kurs_table = fetch_data(db_manager)

    def convert_to_dicts(rows, column_names):
        return [dict(zip(column_names, row)) for row in rows]

    columns_data_table = ['date', 'chk_id', 'dn1_id', 'dn2_id', 'dng_id', 'chnk27d_id', 'pchnk271_id', 'pchnk272_id', 'pchnk273_id', 'schp_id']
    columns_kurs_table = ['id', 'pib']

    kurs_table_dicts = convert_to_dicts(kurs_table, columns_kurs_table)
    kurs_dict = {row['id']: row['pib'] for row in kurs_table_dicts}

    data_table_dicts = convert_to_dicts(data_table, columns_data_table)

    def month_name_genitive(month_number):
        months_genitive = [
            "січня", "лютого", "березня", "квітня", "травня", "червня",
            "липня", "серпня", "вересня", "жовтня", "листопада", "грудня"
        ]
        return months_genitive[month_number - 1]

    data = []
    for row in data_table_dicts:
        day = row['date'].split('-')[2]
        month = month_name_genitive(int(row['date'].split('-')[1]))

        entry = {
            "day": day,
            "month": month,
            'chk': kurs_dict.get(row.get('chk_id'), ''),
            'dn1': kurs_dict.get(row.get('dn1_id'), ''),
            'dn2': kurs_dict.get(row.get('dn2_id'), ''),
            'dng': kurs_dict.get(row.get('dng_id'), ''),
            'chnk27d': kurs_dict.get(row.get('chnk27d_id'), ''),
            'pchnk271': kurs_dict.get(row.get('pchnk271_id'), ''),
            'pchnk272': kurs_dict.get(row.get('pchnk272_id'), ''),
            'pchnk273': kurs_dict.get(row.get('pchnk273_id'), ''),
            'schp': kurs_dict.get(row.get('schp_id'), '')
        }
        data.append(entry)

    for entry in data:
        print(entry)

    template_path = os.path.join(os.path.dirname(__file__), 'doxc.docx')
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"The file '{template_path}' does not exist.")

    def replace_text(doc, replacements):
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            if old_text in paragraph.text:
                                paragraph.text = paragraph.text.replace(old_text, new_text)

    final_doc = Document()

    for entry in data:
        doc = Document(template_path)
        replacements = {f'{key}': value for key, value in entry.items()}
        replace_text(doc, replacements)
        for element in doc.element.body:
            final_doc.element.body.append(deepcopy(element))
        for _ in range(4):  # Add 4 new empty paragraphs
            final_doc.add_paragraph()
        if entry != data[-1]:
            final_doc.add_page_break()

    def remove_empty_paragraphs(doc):
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                p = paragraph._element
                p.getparent().remove(p)
                p._element = None

    def remove_leading_empty_paragraphs(doc):
        while doc.paragraphs and not doc.paragraphs[0].text.strip():
            p = doc.paragraphs[0]._element
            p.getparent().remove(p)
            p._element = None

    def remove_trailing_empty_paragraphs(doc):
        while doc.paragraphs and not doc.paragraphs[-1].text.strip():
            p = doc.paragraphs[-1]._element
            p.getparent().remove(p)
            p._element = None

    def remove_all_empty_until_text(doc):
        # Remove leading empty paragraphs
        remove_leading_empty_paragraphs(doc)
        # Remove trailing empty paragraphs
        remove_trailing_empty_paragraphs(doc)

    remove_all_empty_until_text(final_doc)

    final_doc.save(f'{data[0]["day"]}.{data[0]["month"]}-{data[-1]["day"]}.{data[-1]["month"]}.docx')

# Usage example:
# make_word(db_manager_instance, self_instance)
